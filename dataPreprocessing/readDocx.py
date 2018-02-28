#!/usr/bin/env python
# !coding:utf-8
from databaseInterface.database import *
from aip import AipNlp
import pymysql
import time
import docx


def baiduAPI():
    APP_ID = '10748348'
    API_KEY = 'OtMHyDT0khr8QYP4hjqULKUG'
    SECRET_KEY = 'aGH5pVk3yrdiwX7C4dn5KcyQGhLKdTkj'
    client = AipNlp(APP_ID, API_KEY, SECRET_KEY)
    return client

def sentencesQuestionAnswer(documentName):
    questionAnswer = {}
    maxLenByte = 128
    minLenByte = 10
    # 获取文档对象
    file = docx.Document(documentName)
    # print("段落数:"+str(len(file.paragraphs)))
    client = baiduAPI()
    operating = Database("192.168.160.36", "user_zwb", "123456", "grammer", 3306)
    articleId = operating.selectDataArticleByArticleName(documentName.replace('\\','/').split('/')[-1])
    question = operating.selectDataQuestionAnswerQuestion("!=''")
    for para, paragraphNumber in zip(file.paragraphs, range(len(file.paragraphs))):
        paragraphText = para.text.replace('　', '')
        sentencePeriod = paragraphText.replace('!', '。').replace('；', '。').split('。')
        sentences = []
        for i in sentencePeriod:
            if i != '':
                if len(i) <= minLenByte:
                    continue
                elif len(i) > maxLenByte:
                    sentence = i.split('，')
                    for j in sentence:
                        if j != '':
                            if len(j) < minLenByte:
                                continue
                            else:
                                sentences.append(j)
                else:
                    sentences.append(i)
        # fileOpen.write(str(sentences)+'\n')
        paragraphTextContent = operating.selectDataParagraphText(documentName.replace('\\','/').split('/')[-1])
        if paragraphText not in paragraphTextContent :
            if paragraphText =='':
                continue
            fromParagraph = operating.insertDataQuestionParagraph(articleId,paragraphText)
        else:
            fromParagraph = paragraphTextContent[paragraphText]
        for sentence in sentences:
            flagSQL = True
            try:
                # selectSQL = "SELECT baiduResult FROM baiduResult WHERE sentence=\'%s\'"
                # cur.execute(selectSQL % sentence)
                try:
                    ls = operating.selectDataBaiduResult(sentence)
                    fromProduce = ls[0][1]
                    baidu_result = eval(ls[0][0])
                    flagSQL = False
                except:
                    baidu_result = client.depParser(sentence)
                    try:
                        if baidu_result['error_code'] in [4, 14, 17, 18, 19, 100, 110, 111, 311, 282000, 282002, 282004,
                                                          282130, 282131, 282133, 282300, 282301, 282302, 282303]:
                            continue
                    except Exception as e:
                        pass
            except Exception as e:
                print(e.args)
                continue
            if flagSQL:
                fromProduce = operating.insertDataBaiduResult(articleId,sentence,baidu_result)
                # try:
                #     # sql = '''insert into questionANDanswer(log_id,text,items) VALUES (%d,\"%s\",\"%s\")'''
                #     cur.execute(sql % (baidu_result['log_id'], baidu_result['text'], baidu_result))
                #     cur.lastrowid
                #     connect.commit()
                # except Exception as e:
                #     connect.rollback()
                #     print(e.args)
            whatQuestion = ''
            whoQuestion = ''
            whenQuestion = ''
            howQuestion = ''

            whatNumber = 0
            whoNumber = 0
            whenNumber = 0
            howNumber = 0
            flag = True
            whatFlag = False
            whenFlag = False
            howFlag = False
            whoFlag = False
            for deprel, number in zip(baidu_result['items'], range(len(baidu_result['items']))):
                if flag:
                    if deprel['deprel'] == 'HED':
                        whatQuestion += deprel['word']
                        whatQuestion += '什么'
                        whatNumber = number
                        flag = False
                        whatFlag = True
                    else:
                        whatQuestion += deprel['word']
                if whatFlag:
                    if deprel['deprel'] == 'WP' and number == whatNumber + 1:
                        whatNumber = number
                    elif deprel['deprel'] == 'COO' and number == whatNumber + 1:
                        tmp = whatQuestion[:-2]
                        whatQuestion += tmp + deprel['word'] + '什么'
                        whatNumber = number
                    else:
                        whatFlag = False

                if deprel['deprel'] == 'TMP':
                    whenQuestion += '什么时候'
                    whenFlag = True
                    whenNumber = number
                else:
                    if whenFlag:
                        if (deprel['deprel'] == 'COO' and number == whenNumber + 1) or (
                                deprel['deprel'] == 'WP' and number == whenNumber + 1):
                            whenNumber = number
                        else:
                            whenFlag = False
                    if not whenFlag:
                        whenQuestion += deprel['word']

                if deprel['deprel'] == 'QUN':
                    howQuestion += '多少'
                    howNumber = number
                    howFlag = True
                else:
                    if howFlag:
                        if (deprel['deprel'] == 'COO' and number == howNumber + 1) or (
                                deprel['deprel'] == 'WP' and number == howNumber + 1):
                            howNumber = number
                        else:
                            howFlag = False
                    if not howFlag:
                        howQuestion += deprel['word']

                if deprel['deprel'] == 'SBV':
                    whoQuestion += '谁/什么'
                    whoNumber = number
                    whoFlag = True
                else:
                    if whoFlag:
                        if (deprel['deprel'] == 'COO' and number == whoNumber + 1) or (
                                deprel['deprel'] == 'WP' and number == whoNumber + 1):
                            whoNumber = number
                        else:
                            whoFlag = False
                    if not whoFlag:
                        whoQuestion += deprel['word']

            if '什么' in whatQuestion:
                if whatQuestion+'?' not in question:
                    operating.insertDataQuestionAnswer(articleId,fromParagraph,fromProduce,whatQuestion+'?',[sentence])
                questionAnswer[whatQuestion + '?'] = [sentence]
            if '什么时候' in whenQuestion:
                if whenQuestion+'?' not in question:
                    operating.insertDataQuestionAnswer(articleId,fromParagraph,fromProduce,whenQuestion+'?',[sentence])
                questionAnswer[whenQuestion + '?'] = [sentence]
            if '多少' in howQuestion:
                if howQuestion+'?' not in question:
                    operating.insertDataQuestionAnswer(articleId,fromParagraph,fromProduce,howQuestion+'?',[sentence])
                questionAnswer[howQuestion + '?'] = [sentence]
            if '谁/什么' in whoQuestion:
                if whoQuestion+'?' not in question:
                    operating.insertDataQuestionAnswer(articleId,fromParagraph,fromProduce,whoQuestion+'?',[sentence])
                questionAnswer[whoQuestion + '?'] = [sentence]
    return questionAnswer


def sentencesMain(path):
    return sentencesQuestionAnswer(path)


if __name__ == '__main__':
    print(sentencesMain("../nineteenReportDocuments.docx"))
