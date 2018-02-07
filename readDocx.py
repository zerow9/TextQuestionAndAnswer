#!/usr/bin/env python
#!coding:utf-8
from aip import AipNlp
import pymysql
import time
import docx

def baiduAPI():
    APP_ID = '10748348'
    API_KEY = 'OtMHyDT0khr8QYP4hjqULKUG'
    SECRET_KEY = 'aGH5pVk3yrdiwX7C4dn5KcyQGhLKdTkj'
    client = AipNlp(APP_ID, API_KEY, SECRET_KEY)
    return client

def connectDataBase(host,user,password,database,port):
    try:
        # connect = pymysql.connect(host="192.168.160.36",user="user_zwb",password="123456",db="grammer",port=3306,charset='utf8')
        connect = pymysql.connect(host=host,user=user,password=password,db=database,port=port,charset='utf8')
        cursor = connect.cursor()
        return connect,cursor
    except Exception as e:
        print(e.args)

def sentencesQuestionAnswer(documentName):
    questionAnswer = {}
    Goblenumber = 0
    maxLenByte = 128
    minLenByte = 10
    #获取文档对象
    # file=docx.Document("nineteenReportDocuments.docx")
    file = docx.Document(documentName)
    # print("段落数:"+str(len(file.paragraphs)))
    client = baiduAPI()
    connect,cur = connectDataBase("192.168.160.36","user_zwb","123456","grammer",3306)
    for para,paragraphNumber in zip(file.paragraphs,range(len(file.paragraphs))):
        paragraphText = para.text.replace('　', '')
        sentencePeriod = paragraphText.replace('!','。').replace('；','。').split('。')
        sentences = []
        for i in sentencePeriod:
            if i != '':
                if len(i)<=minLenByte:
                    continue
                elif len(i)>maxLenByte:
                    sentence = i.split('，')
                    for j in sentence:
                        if j != '':
                            if len(j)<minLenByte:
                                continue
                            else:
                                sentences.append(j)
                else:
                    sentences.append(i)
        # fileOpen.write(str(sentences)+'\n')
        for sentence in sentences:
            flagSQL = True
            try:
                selectSQL = "SELECT items FROM questionANDanswer WHERE text=\'%s\'"
                cur.execute(selectSQL % sentence)
                ls = cur.fetchall()
                if len(ls):
                    baidu_result = eval(ls[0][0])
                    flagSQL = False
                else:
                    baidu_result = client.depParser(sentence)
                    # print(baidu_result)
                    try:
                        if baidu_result['error_code'] in [4,14,17,18,19,100,110,111,311,282000,282002,282004,282130,282131,282133,282300,282301,282302,282303]:
                            continue
                    except Exception as e:
                        pass
            except Exception as e:
                print(e.args)
                continue
            if flagSQL:
                try:
                    # sql = '''insert into questionANDanswer(log_id,text,items) VALUES (%d,\"%s\",\"%s\")'''
                    cur.execute(sql % (baidu_result['log_id'], baidu_result['text'], baidu_result))
                    cur.lastrowid
                    connect.commit()
                except Exception as e:
                    connect.rollback()
                    print(e.args)
            # print(sentence)
            # fileOpen.write(str(sentence)+'\n')
            # print(baidu_result)
            # fileOpen.write(str(baidu_result)+'\n')

            whatQuestion = ''
            whoQuestion = ''
            whenQuestion = ''
            howQuestion = ''

            whatNumber = 0
            whoNumber = 0
            whenNumber = 0
            howNumber = 0
            flag = True
            whatFlag = False
            whenFlag = False
            howFlag = False
            whoFlag = False
            for deprel,number in zip(baidu_result['items'],range(len(baidu_result['items']))):
                if flag:
                    if deprel['deprel'] == 'HED':
                        whatQuestion += deprel['word']
                        whatQuestion += '什么'
                        whatNumber = number
                        flag = False
                        whatFlag = True
                    else:
                        whatQuestion += deprel['word']
                if whatFlag:
                    if deprel['deprel'] == 'WP'and number == whatNumber+1:
                        whatNumber = number
                    elif deprel['deprel'] == 'COO'and number == whatNumber+1:
                        tmp = whatQuestion[:-2]
                        whatQuestion += tmp + deprel['word'] + '什么'
                        whatNumber = number
                    else:
                        whatFlag = False

                if deprel['deprel'] == 'TMP':
                    whenQuestion += '什么时候'
                    whenFlag = True
                    whenNumber = number
                else:
                    if whenFlag:
                        if (deprel['deprel'] == 'COO' and number == whenNumber + 1) or (deprel['deprel'] == 'WP' and number == whenNumber + 1):
                            whenNumber = number
                        else:
                            whenFlag = False
                    if not whenFlag:
                        whenQuestion += deprel['word']

                if deprel['deprel'] == 'QUN':
                    howQuestion += '多少'
                    howNumber = number
                    howFlag = True
                else:
                    if howFlag:
                        if (deprel['deprel'] == 'COO' and number == howNumber + 1) or (deprel['deprel'] == 'WP' and number == howNumber + 1):
                            howNumber = number
                        else:
                            howFlag = False
                    if not howFlag:
                        howQuestion += deprel['word']

                if deprel['deprel'] == 'SBV':
                    whoQuestion += '谁/什么'
                    whoNumber = number
                    whoFlag = True
                else:
                    if whoFlag:
                        if (deprel['deprel'] == 'COO' and number == whoNumber + 1) or (deprel['deprel'] == 'WP' and number == whoNumber + 1):
                            whoNumber = number
                        else:
                            whoFlag = False
                    if not whoFlag:
                        whoQuestion += deprel['word']

            # FrequentlyAskedQuestions.write('Question:\n')
            if '什么' in whatQuestion:
                questionAnswer[whatQuestion+'?'] = sentence
                # FrequentlyAskedQuestions.write('\t'+whatQuestion+'?\n')
            if '什么时候' in whenQuestion:
                questionAnswer[whenQuestion+'?'] = sentence
                # FrequentlyAskedQuestions.write('\t'+whenQuestion+'?\n')
            if '多少' in howQuestion:
                questionAnswer[howQuestion+'?'] = sentence
                # FrequentlyAskedQuestions.write('\t'+howQuestion+'?\n')
            if '谁/什么' in whoQuestion:
                questionAnswer[whoQuestion+'?'] = sentence
                # FrequentlyAskedQuestions.write('\t'+whoQuestion+'?\n')
            # FrequentlyAskedQuestions.write('Answer:\n')
            # FrequentlyAskedQuestions.write('\t'+sentence+'\n\n')
            Goblenumber += 1
            # print(Goblenumber)
            # fileOpen.write('\n')
    connect.close()
    return questionAnswer

def sentencesMain(path):
    # fileOpen = open('testXXXX.txt', 'w', encoding='utf-8')
    # FrequentlyAskedQuestions = open('FrequentlyAskedQuestionsQueryXXXXXX.txt', 'w', encoding='utf-8')
    return sentencesQuestionAnswer(path)
    # fileOpen.close()
    # FrequentlyAskedQuestions.close()

if __name__ == '__main__':
    print(sentencesMain("nineteenReportDocuments.docx"))